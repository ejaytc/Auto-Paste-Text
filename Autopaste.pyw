#!/usr/bin/env python3 

from pynput.keyboard import Key, Listener, KeyCode
from docx import Document
import os, pyperclip, shelve, subprocess, re


def autopaste():
    data = shelve.open("AutopasteData")
    Path = data["path_"]
    File = data["file_"]
    data.close()
    check_(Path)
    
    new_document = os.path.join(Path, File)
    if not os.path.exists(new_document):
        document = Document()
    else:
        document = Document(new_document)
    
    text         = pyperclip.paste()
    document.add_paragraph(text)
    document.save(new_document)

def open_ui():
    check_(file_path)
    subprocess.call("python3 ApMainUi.pyw", shell=True)


def check_(file_path):
    if not os.path.exists(file_path):
        os.mkdir(file_path)

COMBINATION = {
    frozenset([Key.shift, KeyCode(char = 'S')]): autopaste,
    frozenset([Key.shift, KeyCode(char = 's')]): autopaste,
    frozenset([Key.ctrl, Key.alt, KeyCode(char='U')]): open_ui,
    frozenset([Key.ctrl, Key.alt, KeyCode(char='u')]): open_ui
}

current = set()
def on_press(key):
    current.add(key)
    if frozenset(current) in COMBINATION:
        COMBINATION[frozenset(current)]()

def on_released(key):
    if key in current:
        current.remove(key)
    elif not current == None:
        rm = list(current)
        for r in rm:
            current.remove(r)
    
if __name__ == "__main__":
    fformat     = re.compile(r'\D*.docx')
    data        = shelve.open("AutopasteData")

    if not 'True' in dict(data):
        data["path_"] = "/home/jay/Documents/AutoPasteDocuments"
        data["file_"] = "pastedDoc-file.docx"
        data["Check"] = True

    file_path = data["path_"]
    data.close()
    check_(file_path)
    with Listener(on_press=on_press, on_release=on_released) as listener:
        listener.join()